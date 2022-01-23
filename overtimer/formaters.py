from docxtpl import DocxTemplate
import docx
import os
import datetime
import random
import string
from copy import deepcopy

TEMPLATE_PATH = 'templates/template.docx'


def get_tempfile_name(temp_dir_path):
    letters = string.ascii_lowercase
    rand_string = ''.join(random.choice(letters) for i in range(6))
    return os.path.join(temp_dir_path, rand_string)


def make_docx(contexts, folder=os.getcwd()):
    templ_filenames = []
    today = datetime.datetime.today()
    core_filename = '{}.docx'.format(today.strftime("%d%b%Y"))
    for context in contexts:
        document = DocxTemplate(TEMPLATE_PATH)
        document.render(context)
        temp_name = get_tempfile_name(folder)
        document.save(temp_name)
        templ_filenames.append(temp_name)
    core_templete_file_name = templ_filenames.pop(0)
    core_doc = docx.Document(core_templete_file_name)
    os.remove(core_templete_file_name)
    for filename in templ_filenames:
        paragraphs = docx.Document(filename)
        tbl = paragraphs.tables[0]._tbl
        new_table = deepcopy(tbl)
        core_paragraph = core_doc.add_paragraph()
        core_paragraph._p.addnext(new_table)
        os.remove(filename)
    core_doc.save(os.path.join(folder, core_filename))
    print(os.path.join(folder, core_filename))
